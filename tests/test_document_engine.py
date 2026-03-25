"""
文件處理引擎 & 進階檢索 — 單元測試套件

測試範圍：
  1. DocumentParser — 多格式解析 + 品質報告
  2. TextChunker — 精確 Token 切片 + 章節偵測
  3. KnowledgeBaseRetriever — 進階檢索功能驗證
"""

import os
import sys
import json
import tempfile
import time

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# ═══════════════════════════════════════════════
# Colored output helpers
# ═══════════════════════════════════════════════
GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
CYAN = "\033[96m"
BOLD = "\033[1m"
RESET = "\033[0m"

passed = 0
failed = 0
warnings = 0


def test(name: str, condition: bool, detail: str = ""):
    global passed, failed
    if condition:
        passed += 1
        print(f"  {GREEN}✓{RESET} {name}")
    else:
        failed += 1
        print(f"  {RED}✗{RESET} {name}")
        if detail:
            print(f"    {RED}→ {detail}{RESET}")


def warn(msg: str):
    global warnings
    warnings += 1
    print(f"  {YELLOW}⚠{RESET} {msg}")


def section(title: str):
    print(f"\n{BOLD}{CYAN}{'=' * 60}{RESET}")
    print(f"{BOLD}{CYAN}{title}{RESET}")
    print(f"{BOLD}{CYAN}{'=' * 60}{RESET}\n")


# ═══════════════════════════════════════════════
# 1. DocumentParser 測試
# ═══════════════════════════════════════════════

section("1. DocumentParser — 格式支援與偵測")

from app.services.document_parser import (
    DocumentParser,
    TextChunker,
    QualityReport,
    SUPPORTED_FORMATS,
    _HAS_TIKTOKEN,
    _HAS_PDFPLUMBER,
    _HAS_OPENPYXL,
    _HAS_OCR,
    _HAS_CHARDET,
    _HAS_RTF,
)

# 1.1 格式映射完整性
test("SUPPORTED_FORMATS 包含 PDF",    ".pdf" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 DOCX",   ".docx" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 DOC",    ".doc" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 TXT",    ".txt" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 XLSX",   ".xlsx" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 XLS",    ".xls" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 CSV",    ".csv" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 HTML",   ".html" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 HTM",    ".htm" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 MD",     ".md" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 RTF",    ".rtf" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 JSON",   ".json" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 JPG",    ".jpg" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 PNG",    ".png" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 TIFF",   ".tiff" in SUPPORTED_FORMATS)
test("SUPPORTED_FORMATS 包含 BMP",    ".bmp" in SUPPORTED_FORMATS)
test("支援格式總數 >= 16",            len(SUPPORTED_FORMATS) >= 16)

# 1.2 detect_file_type
test("detect_file_type('.pdf') = 'pdf'",    DocumentParser.detect_file_type("test.pdf") == "pdf")
test("detect_file_type('.xlsx') = 'xlsx'",  DocumentParser.detect_file_type("data.xlsx") == "xlsx")
test("detect_file_type('.html') = 'html'",  DocumentParser.detect_file_type("page.html") == "html")
test("detect_file_type('.csv') = 'csv'",    DocumentParser.detect_file_type("data.csv") == "csv")
test("detect_file_type('.md') = 'markdown'", DocumentParser.detect_file_type("README.md") == "markdown")
test("detect_file_type('.json') = 'json'",  DocumentParser.detect_file_type("config.json") == "json")
test("detect_file_type('.jpg') = 'image'",  DocumentParser.detect_file_type("photo.jpg") == "image")

try:
    DocumentParser.detect_file_type("test.xyz")
    test("不支援格式拋 ValueError", False)
except ValueError:
    test("不支援格式拋 ValueError", True)

# 1.3 依賴偵測
print(f"\n  {CYAN}依賴偵測狀態：{RESET}")
print(f"    tiktoken:     {'✓ 已安裝' if _HAS_TIKTOKEN else '✗ 未安裝'}")
print(f"    pdfplumber:   {'✓ 已安裝' if _HAS_PDFPLUMBER else '✗ 未安裝'}")
print(f"    openpyxl:     {'✓ 已安裝' if _HAS_OPENPYXL else '✗ 未安裝'}")
print(f"    chardet:      {'✓ 已安裝' if _HAS_CHARDET else '✗ 未安裝'}")
print(f"    striprtf:     {'✓ 已安裝' if _HAS_RTF else '✗ 未安裝'}")
print(f"    pytesseract:  {'✓ 已安裝' if _HAS_OCR else '✗ 未安裝'}")

# ═══════════════════════════════════════════════
# 2. 實際文件解析測試（使用臨時文件）
# ═══════════════════════════════════════════════

section("2. DocumentParser — 實際解析測試")

# 2.1 TXT 解析
with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", encoding="utf-8", delete=False) as f:
    f.write("這是一段測試文字。\n\n員工請假辦法如下：\n\n第一條、適用範圍：本辦法適用於全體正式員工。\n第二條、請假種類：病假、事假、特休假、婚假、喪假、產假。\n第三條、病假規定：每年病假不超過三十日，以半薪計算。")
    txt_path = f.name

text, meta = DocumentParser.parse(txt_path, "txt")
test("TXT 解析: 文字內容長度 > 50",     len(text) > 50)
test("TXT 解析: quality_level 不為 failed", meta.get("quality_level") != "failed")
test("TXT 解析: 包含 '請假'",            "請假" in text)
test("TXT 解析: parse_time_ms 存在",      "parse_time_ms" in meta)
os.unlink(txt_path)

# 2.2 Markdown 解析
with tempfile.NamedTemporaryFile(suffix=".md", mode="w", encoding="utf-8", delete=False) as f:
    f.write("# 員工手冊\n\n## 第一章 總則\n\n本手冊適用於全體員工。\n\n## 第二章 出勤\n\n上午九點上班，下午六點下班。\n\n### 2.1 彈性工時\n\n可申請彈性工時。")
    md_path = f.name

text, meta = DocumentParser.parse(md_path, "markdown")
test("MD 解析: 包含標題 '員工手冊'",     "員工手冊" in text)
test("MD 解析: format_detected = 'markdown'", meta.get("format_detected") == "markdown")
os.unlink(md_path)

# 2.3 CSV 解析
with tempfile.NamedTemporaryFile(suffix=".csv", mode="w", encoding="utf-8", delete=False) as f:
    f.write("姓名,部門,假別,天數\n張三,工程部,特休,7\n李四,業務部,病假,3\n王五,人資部,事假,2")
    csv_path = f.name

text, meta = DocumentParser.parse(csv_path, "csv")
test("CSV 解析: tables_detected = 1",    meta.get("tables_detected") == 1)
test("CSV 解析: 包含 '張三'",           "張三" in text)
test("CSV 解析: 包含 '工程部'",          "工程部" in text)
os.unlink(csv_path)

# 2.4 HTML 解析
with tempfile.NamedTemporaryFile(suffix=".html", mode="w", encoding="utf-8", delete=False) as f:
    f.write("""<html><head><title>公司規定</title><style>body{}</style></head>
<body><h1>工作規則</h1><p>本規則適用於全體員工。</p>
<h2>出勤管理</h2><p>上午九點上班。</p>
<table><tr><th>假別</th><th>天數</th></tr><tr><td>特休</td><td>7</td></tr></table>
<script>alert('test')</script></body></html>""")
    html_path = f.name

text, meta = DocumentParser.parse(html_path, "html")
test("HTML 解析: 包含 '工作規則'",        "工作規則" in text)
test("HTML 解析: 不包含 script 內容",     "alert" not in text)
test("HTML 解析: tables_detected >= 1",   meta.get("tables_detected", 0) >= 1)
os.unlink(html_path)

# 2.5 JSON 解析
with tempfile.NamedTemporaryFile(suffix=".json", mode="w", encoding="utf-8", delete=False) as f:
    json.dump({
        "company": "測試公司",
        "policies": [
            {"name": "請假辦法", "content": "病假三十日"},
            {"name": "加班辦法", "content": "平日加班費 1.34 倍"}
        ]
    }, f, ensure_ascii=False)
    json_path = f.name

text, meta = DocumentParser.parse(json_path, "json")
test("JSON 解析: 包含 '測試公司'",        "測試公司" in text)
test("JSON 解析: 包含 '病假三十日'",       "病假三十日" in text)
os.unlink(json_path)

# 2.6 Excel 解析
if _HAS_OPENPYXL:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "薪資表"
    ws.append(["職級", "基本薪資", "加班費率"])
    ws.append(["工程師", "50000", "1.34"])
    ws.append(["資深工程師", "65000", "1.34"])
    ws.append(["主管", "80000", "1.67"])
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        xlsx_path = f.name
    wb.save(xlsx_path)

    text, meta = DocumentParser.parse(xlsx_path, "xlsx")
    test("Excel 解析: tables_detected >= 1",  meta.get("tables_detected", 0) >= 1)
    test("Excel 解析: 包含 '工程師'",         "工程師" in text)
    test("Excel 解析: 包含 '50000'",          "50000" in text)
    os.unlink(xlsx_path)
else:
    warn("openpyxl 未安裝，跳過 Excel 解析測試")

# 2.7 RTF 解析
if _HAS_RTF:
    with tempfile.NamedTemporaryFile(suffix=".rtf", mode="w", encoding="utf-8", delete=False) as f:
        f.write(r"{\rtf1\ansi\deff0{\fonttbl{\f0 Times New Roman;}}{\pard This is a test document about leave policy.\par}}")
        rtf_path = f.name

    text, meta = DocumentParser.parse(rtf_path, "rtf")
    test("RTF 解析: 包含 'leave policy'",  "leave policy" in text.lower())
    os.unlink(rtf_path)
else:
    warn("striprtf 未安裝，跳過 RTF 解析測試")

# 2.8 DOCX 解析
from docx import Document as DocxDoc
doc = DocxDoc()
doc.add_heading("公司工作規則", level=1)
doc.add_paragraph("本規則適用於本公司全體員工。")
doc.add_heading("第一章 薪資", level=2)
doc.add_paragraph("基本薪資不得低於法定最低工資。")
table = doc.add_table(rows=3, cols=2)
table.cell(0, 0).text = "項目"
table.cell(0, 1).text = "金額"
table.cell(1, 0).text = "基本薪資"
table.cell(1, 1).text = "27470"
table.cell(2, 0).text = "加班費"
table.cell(2, 1).text = "依法計算"
with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    docx_path = f.name
doc.save(docx_path)

text, meta = DocumentParser.parse(docx_path, "docx")
test("DOCX 解析: 包含標題 '工作規則'",     "工作規則" in text)
test("DOCX 解析: tables_detected >= 1",    meta.get("tables_detected", 0) >= 1)
test("DOCX 解析: 包含表格內容 '27470'",    "27470" in text)
test("DOCX 解析: 偵測到標題層級 '#'",       "#" in text)
os.unlink(docx_path)


# ═══════════════════════════════════════════════
# 3. QualityReport 測試
# ═══════════════════════════════════════════════

section("3. QualityReport — 品質報告系統")

report = QualityReport(format_detected="pdf", total_chars=5000, total_pages=10)
report.compute_quality()
test("品質分數 = 1.0 (無問題)", report.quality_score == 1.0)
test("品質等級 = excellent",    report.quality_level == "excellent")

report2 = QualityReport(format_detected="pdf", total_chars=50)
report2.add_warning("第 3 頁為掃描")
report2.add_warning("表格格式可能遺失")
report2.compute_quality()
test("有 2 警告: quality_score < 1.0", report2.quality_score < 1.0)
test("品質等級不為 excellent",         report2.quality_level != "excellent")

report3 = QualityReport(format_detected="pdf", total_chars=10)
report3.add_error("OCR 失敗")
report3.add_error("無法提取文字")
report3.compute_quality()
test("有錯誤+低字數: quality_level = poor 或 failed", report3.quality_level in ("poor", "failed"))

d = report.to_dict()
test("to_dict() 包含 quality_score",   "quality_score" in d)
test("to_dict() 包含 format_detected",  "format_detected" in d)
test("to_dict() 包含 warnings",         "warnings" in d)


# ═══════════════════════════════════════════════
# 4. TextChunker 測試
# ═══════════════════════════════════════════════

section("4. TextChunker — 智慧切片")

# 4.1 Token 計算
token_count = TextChunker.count_tokens("Hello world. 你好世界。")
test(f"Token 計算: 'Hello world. 你好世界。' = {token_count} tokens", token_count > 0)

if _HAS_TIKTOKEN:
    test("使用 tiktoken 精確計算", True)
else:
    warn("tiktoken 未安裝，使用估算模式")

# 4.2 基本切片
long_text = "\n\n".join([
    f"第{i}條 這是一段關於員工管理的測試文字，包含各種規定和細節說明。" * 5
    for i in range(1, 51)
])
chunks = TextChunker.split_by_tokens(long_text, chunk_size=200, chunk_overlap=30)
test(f"長文本切片: 產生 {len(chunks)} 個 chunks > 0",  len(chunks) > 0)
test("每個 chunk 不為空",  all(len(c.strip()) > 0 for c in chunks))

# 4.3 表格保護
table_text = """# 薪資規定

基本薪資規定如下：

[表格 1]
職級 | 薪資 | 加班費率
工程師 | 50000 | 1.34
資深工程師 | 65000 | 1.34
主管 | 80000 | 1.67

以上表格為公司薪資標準。

# 請假規定

員工請假需提前申請。"""

table_chunks = TextChunker.split_by_tokens(table_text, chunk_size=500, chunk_overlap=50)
# 檢查表格是否被保護
table_intact = any("[表格 1]" in c and "主管" in c for c in table_chunks)
test("表格保護: 表格內容在同一 chunk 中", table_intact)

# 4.4 標題邊界切分
heading_text = """# 第一章 總則

本章說明基本規定。

# 第二章 出勤

上午九點上班。

# 第三章 請假

請假需提前申請。"""

heading_chunks = TextChunker.split_by_tokens(heading_text, chunk_size=500, chunk_overlap=50)
test(f"標題切分: 產生 {len(heading_chunks)} 個 chunks", len(heading_chunks) >= 1)

# 4.5 空文本
empty_chunks = TextChunker.split_by_tokens("", chunk_size=500)
test("空文本: 返回 []", empty_chunks == [])

# 4.6 很短的文本（低於最低門檻）
short_chunks = TextChunker.split_by_tokens("短", chunk_size=500)
test("過短文本: 被過濾", short_chunks == [])


# ═══════════════════════════════════════════════
# 5. KnowledgeBaseRetriever 結構測試
# ═══════════════════════════════════════════════

section("5. KnowledgeBaseRetriever — 結構驗證")

from app.services.kb_retrieval import KnowledgeBaseRetriever, _HAS_BM25

test("BM25 library 已安裝", _HAS_BM25)

# 驗證類的方法
test("有 search 方法",            hasattr(KnowledgeBaseRetriever, "search"))
test("有 batch_search 方法",      hasattr(KnowledgeBaseRetriever, "batch_search"))
test("有 get_stats 方法",         hasattr(KnowledgeBaseRetriever, "get_stats"))
test("有 _semantic_search 方法",  hasattr(KnowledgeBaseRetriever, "_semantic_search"))
test("有 _keyword_search 方法",   hasattr(KnowledgeBaseRetriever, "_keyword_search"))
test("有 _hybrid_search 方法",    hasattr(KnowledgeBaseRetriever, "_hybrid_search"))
test("有 _rerank 方法",           hasattr(KnowledgeBaseRetriever, "_rerank"))
test("有 invalidate_cache 方法",  hasattr(KnowledgeBaseRetriever, "invalidate_cache"))

# 驗證 search 方法簽名
import inspect
sig = inspect.signature(KnowledgeBaseRetriever.search)
params = list(sig.parameters.keys())
test("search() 有 mode 參數",      "mode" in params)
test("search() 有 min_score 參數",  "min_score" in params)
test("search() 有 rerank 參數",     "rerank" in params)
test("search() 有 use_cache 參數",  "use_cache" in params)

# BM25 分詞測試
tokenize = KnowledgeBaseRetriever._tokenize
tokens = tokenize("員工請假辦法 employee leave policy")
test(f"中英混合分詞: 產生 {len(tokens)} 個 tokens > 5", len(tokens) > 5)
test("包含中文字元",  "員" in tokens)
test("包含英文詞",    "employee" in tokens)


# ═══════════════════════════════════════════════
# 6. 設定檢查
# ═══════════════════════════════════════════════

section("6. 設定 & Schema 完整性")

from app.config import settings

test("settings 有 RETRIEVAL_MODE",       hasattr(settings, "RETRIEVAL_MODE"))
test("settings 有 RETRIEVAL_MIN_SCORE",  hasattr(settings, "RETRIEVAL_MIN_SCORE"))
test("settings 有 RETRIEVAL_RERANK",     hasattr(settings, "RETRIEVAL_RERANK"))
test("settings 有 RETRIEVAL_CACHE_TTL",  hasattr(settings, "RETRIEVAL_CACHE_TTL"))
test("settings 有 RETRIEVAL_TOP_K",      hasattr(settings, "RETRIEVAL_TOP_K"))
test("RETRIEVAL_MODE = 'hybrid'",        settings.RETRIEVAL_MODE == "hybrid")

from app.schemas.document import DocumentUpdate

du = DocumentUpdate(quality_report={"quality_score": 0.9, "warnings": []})
test("DocumentUpdate 支援 quality_report", du.quality_report is not None)

from app.models.document import DocumentChunk
test("DocumentChunk 有 vector_id 欄位", hasattr(DocumentChunk, "vector_id"))

# ═══════════════════════════════════════════════
# 7. 效能測試
# ═══════════════════════════════════════════════

section("7. 效能基準測試")

# 大文件切片效能
big_text = ("這是一段很長的測試文字，用來模擬大型企業文件的內容。" * 100 + "\n\n") * 50
start_time = time.time()
big_chunks = TextChunker.split_by_tokens(big_text, chunk_size=1000, chunk_overlap=150)
elapsed = (time.time() - start_time) * 1000
test(f"大文件切片 ({len(big_text)} 字): {elapsed:.0f}ms < 5000ms", elapsed < 5000)
test(f"大文件切片: 產生 {len(big_chunks)} 個 chunks", len(big_chunks) > 0)

# TXT 解析效能
with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", encoding="utf-8", delete=False) as f:
    f.write(big_text)
    big_txt_path = f.name
start_time = time.time()
text, meta = DocumentParser.parse(big_txt_path, "txt")
parse_elapsed = (time.time() - start_time) * 1000
test(f"大 TXT 解析: {parse_elapsed:.0f}ms < 3000ms", parse_elapsed < 3000)
os.unlink(big_txt_path)


# ═══════════════════════════════════════════════
# 總結
# ═══════════════════════════════════════════════

section("測試總結")

total = passed + failed
print(f"\n  通過: {GREEN}{passed}{RESET}")
print(f"  失敗: {RED}{failed}{RESET}")
print(f"  警告: {YELLOW}{warnings}{RESET}")
print(f"  通過率: {GREEN if failed == 0 else YELLOW}{passed}/{total} ({passed/total*100:.1f}%){RESET}")
print()

if failed == 0:
    print(f"  {GREEN}{BOLD}╔═══════════════════════════════════════════════════╗{RESET}")
    print(f"  {GREEN}{BOLD}║            ✓ ALL TESTS PASSED                    ║{RESET}")
    print(f"  {GREEN}{BOLD}╚═══════════════════════════════════════════════════╝{RESET}")
else:
    print(f"  {RED}{BOLD}╔═══════════════════════════════════════════════════╗{RESET}")
    print(f"  {RED}{BOLD}║         ✗ {failed} TEST(S) FAILED                     ║{RESET}")
    print(f"  {RED}{BOLD}╚═══════════════════════════════════════════════════╝{RESET}")

print(f"""
{CYAN}新增能力摘要：{RESET}
  📄 文件格式: {len(SUPPORTED_FORMATS)} 種（PDF/DOCX/DOC/TXT/Excel/CSV/HTML/MD/RTF/JSON/圖片）
  🔢 Token 計算: {'tiktoken 精確計算' if _HAS_TIKTOKEN else '估算模式'}
  📊 PDF 表格: {'pdfplumber 啟用' if _HAS_PDFPLUMBER else '未啟用'}
  📋 Excel: {'openpyxl 啟用' if _HAS_OPENPYXL else '未啟用'}
  🔍 OCR: {'pytesseract 啟用' if _HAS_OCR else '未安裝（需系統安裝 tesseract）'}
  🔠 編碼偵測: {'chardet 啟用' if _HAS_CHARDET else '未啟用'}
  📑 RTF: {'striprtf 啟用' if _HAS_RTF else '未啟用'}
  🔎 BM25 關鍵字檢索: {'rank-bm25 啟用' if _HAS_BM25 else '未啟用'}
  🔄 混合檢索 (RRF): 已實現
  ⚡ 重排序 (Voyage Rerank): 已實現
  🗄️ Redis 查詢快取: 已實現
""")

if __name__ == "__main__":
    sys.exit(0 if failed == 0 else 1)
