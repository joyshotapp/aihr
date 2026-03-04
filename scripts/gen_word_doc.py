from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm

def set_chinese_font(run, size=12, bold=False, color=None):
    run.font.name = 'Microsoft JhengHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1, color=RGBColor(26, 86, 219)):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_chinese_font(run, size=16 if level==1 else 14, bold=True, color=color)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(12)
    return h

def add_paragraph(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.line_spacing = 1.5
    return p

def create_document():
    doc = Document()
    
    # --- Cover Page ---
    # Removed unnecessary add_section() which caused a blank first page
    
    # Logo / Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(150)
    run = p_title.add_run("UniAI人資長")
    set_chinese_font(run, size=42, bold=True, color=RGBColor(26, 86, 219))
    
    # Tagline
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tag.add_run("讓每一位員工的提問，都有人資長等級的即時回覆")
    set_chinese_font(run, size=16, color=RGBColor(68, 68, 68))
    
    # Version
    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ver.paragraph_format.space_before = Pt(200)
    run = p_ver.add_run("產品簡介 v1.0 ｜ 2026年2月")
    set_chinese_font(run, size=10, color=RGBColor(153, 153, 153))
    
    doc.add_page_break()
    
    # --- Page 1: Challenges & Intro ---
    add_heading(doc, "您的HR部門，是否也面臨這些挑戰？")
    
    challenges = [
        "員工反覆詢問相同的請假規定、加班費算法、福利制度，HR疲於回應",
        "公司內規與勞基法是否衝突，卻沒有人力逐條比對",
        "薪資計算、資遣費試算仰賴人工，既耗時又怕算錯",
        "員工手冊、契約書、管理辦法散落各處，查詢費時費力",
        "想請教法規問題，但聘請法律顧問成本太高"
    ]
    
    for ch in challenges:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(ch)
        set_chinese_font(run, size=11)
        
    add_paragraph(doc, "\n對中小企業而言，HR通常身兼數職，不可能每天花兩三個小時處理重複性問答。UniAI人資長正是為此而生。")
    
    add_heading(doc, "UniAI人資長是什麼？")
    add_paragraph(doc, "一套專為台灣中小企業設計的人力資源智能問答系統。")
    add_paragraph(doc, "只要將公司既有的HR文件上傳，系統即可自動解析內容，結合台灣現行勞動法規，讓員工和HR透過自然語言提問，即時獲得專業回覆。")
    
    p_highlight = doc.add_paragraph()
    p_highlight.paragraph_format.left_indent = Cm(1)
    run = p_highlight.add_run("不需要額外整理資料，不需要標記關鍵字，上傳就能用。")
    set_chinese_font(run, size=11, bold=True, color=RGBColor(22, 163, 74)) # Greenish

    # --- Core Capabilities ---
    add_heading(doc, "三大核心能力")
    
    caps = [
        ("1. 上傳文件，自動變成知識庫", "支援 PDF、Excel、Word、JPG (OCR) 等常見格式。無論是員工手冊、薪資條、勞動契約或掃描文件，上傳後系統自動解析索引，平均每份文件處理時間不到10秒。"),
        ("2. 結合內規與法令，精準回覆", "系統內建台灣勞動法規知識庫。回覆時自動比對「公司內規」與「法令」，若發現衝突會主動提醒。不只答「公司規定」，更答「合不合法」。"),
        ("3. 跨文件推理與數據計算", "讀取員工名冊、薪資表等結構化資料。直接提問「技術部平均薪資」、「特休剩幾天」，系統即時計算。HR不必再開Excel拉公式。")
    ]
    
    for title, desc in caps:
        p = doc.add_paragraph()
        run_t = p.add_run(title + "\n")
        set_chinese_font(run_t, size=12, bold=True, color=RGBColor(30, 64, 175))
        run_d = p.add_run(desc)
        set_chinese_font(run_d, size=11)
        p.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # --- Page 2: Cases ---
    add_heading(doc, "實際導入案例")
    add_paragraph(doc, "以下為台灣中小企業實際導入測試成果。上傳11份文件，驗證43道問題，準確率達到100%。")
    
    # Case 1
    add_heading(doc, "案例一：員工自助查詢福利與假勤", level=2)
    
    qa_list = [
        ("員工：我們公司有交通津貼嗎？補助多少？", "UniAI：公司提供交通津貼，依通勤距離補貼 500-2,000 元。金額範圍依員工手冊規定辦理。"),
        ("員工：員工年資 3 年 2 個月，特休有幾天？", "UniAI：年資約 3.17 年，特休為 14 天。")
    ]
    
    for q, a in qa_list:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.left_indent = Cm(0.5)
        run_q = p_q.add_run(q)
        set_chinese_font(run_q, size=10.5, bold=True, color=RGBColor(80, 80, 80))
        
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Cm(0.5)
        p_a.paragraph_format.space_after = Pt(8)
        run_a = p_a.add_run(a)
        set_chinese_font(run_a, size=10.5, color=RGBColor(26, 86, 219))
        
    # Case 2
    add_heading(doc, "案例二：自動偵測公司規定是否違反勞基法", level=2)
    
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Cm(0.5)
    run_q = p_q.add_run("員工：我們公司平日加班給 1.5 倍工資，這樣合法嗎？")
    set_chinese_font(run_q, size=10.5, bold=True, color=RGBColor(80, 80, 80))
    
    p_a = doc.add_paragraph()
    p_a.paragraph_format.left_indent = Cm(0.5)
    run_a = p_a.add_run("UniAI：公司平日加班給 1.5 倍，前 2 小時優於法定 (1.34倍) 合法；但第 3 小時起法定需 1.67 倍，若僅給 1.5 倍則不足。")
    set_chinese_font(run_a, size=10.5, color=RGBColor(26, 86, 219))

    # Case 2 Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['時段', '公司倍率', '法定倍率', '判定結果']
    for i, h in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        set_chinese_font(run, size=10, bold=True)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    row = table.add_row().cells
    data = [('前2小時', '1.5 倍', '1.34 倍', '合法 (優於法令)'), ('第3小時起', '1.5 倍', '1.67 倍', '未達法定標準')]
    
    for r_data in data:
        row = table.add_row().cells
        for i, txt in enumerate(r_data):
            run = row[i].paragraphs[0].add_run(txt)
            set_chinese_font(run, size=10)
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # Case 3
    add_heading(doc, "案例三：即時計算資遣費、薪資統計", level=2)
    
    qa_list_3 = [
        ("HR：如果今天要資遣 E007 劉志明，需要付多少資遣費？", "UniAI：E007 劉志明年資 6.2 年 (進位 6.5)。資遣費 = 6.5 x 0.5 x 52,000 = 169,000 元 (新制)。"),
        ("HR：技術部的平均月薪是多少？", "UniAI：技術部平均月薪約 59,833 元 (共 6 人平均)。")
    ]
    for q, a in qa_list_3:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.left_indent = Cm(0.5)
        run_q = p_q.add_run(q)
        set_chinese_font(run_q, size=10.5, bold=True, color=RGBColor(80, 80, 80))
        
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Cm(0.5)
        p_a.paragraph_format.space_after = Pt(8)
        run_a = p_a.add_run(a)
        set_chinese_font(run_a, size=10.5, color=RGBColor(26, 86, 219))

    doc.add_page_break()

    # --- Page 3: Comparison & Pricing ---
    add_heading(doc, "為什麼選擇UniAI人資長？")
    
    table_comp = doc.add_table(rows=1, cols=2)
    table_comp.style = 'Table Grid'
    table_comp.rows[0].cells[0].paragraphs[0].add_run("傳統做法").bold = True
    table_comp.rows[0].cells[1].paragraphs[0].add_run("導入UniAI人資長").bold = True
    
    comp_data = [
        ("員工問HR，HR翻手冊，平均3-5分鐘", "員工直接問系統，平均1秒"),
        ("法規合規靠法務或外部顧問，成本高", "系統自動比對，即時識別風險"),
        ("薪資、資遣費手動算，容易出錯", "輸入條件自動計算，公式透明"),
        ("人員異動導致知識斷層", "知識庫永久留存，經驗不流失")
    ]
    
    for old, new in comp_data:
        row = table_comp.add_row().cells
        set_chinese_font(row[0].paragraphs[0].add_run(old), size=10)
        run_new = row[1].paragraphs[0].add_run(new)
        set_chinese_font(run_new, size=10, bold=True, color=RGBColor(26, 86, 219))

    add_heading(doc, "授權方案")
    
    table_price = doc.add_table(rows=1, cols=3)
    table_price.style = 'Table Grid'
    headers_price = ["方案", "適用規模", "功能"]
    for i, h in enumerate(headers_price):
        run = table_price.rows[0].cells[i].paragraphs[0].add_run(h)
        set_chinese_font(run, size=10, bold=True)
        
    price_data = [
        ("基礎版", "50人以下", "文檔解析、基礎問答、假勤查詢"),
        ("專業版", "50-200人", "+ 合規檢測、數據分析、多輪對話"),
        ("企業版", "200人以上", "+ API串接、客製化、專屬部署")
    ]
    
    for p_name, p_size, p_feat in price_data:
        row = table_price.add_row().cells
        set_chinese_font(row[0].paragraphs[0].add_run(p_name), size=10, bold=True)
        set_chinese_font(row[1].paragraphs[0].add_run(p_size), size=10)
        set_chinese_font(row[2].paragraphs[0].add_run(p_feat), size=10)
        
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note.paragraph_format.space_before = Pt(6)
    run_note = p_note.add_run("* 所有方案均提供 30 天免費試用，導入前可充分驗證效果。")
    set_chinese_font(run_note, size=9, color=RGBColor(100, 100, 100))
    
    add_heading(doc, "立即體驗")
    
    cta_text = [
        "預約線上Demo：實際操作給您看",
        "免費試用30天：上傳貴公司文件，親自驗證效果",
        "客製化諮詢：針對特殊需求提供專屬方案"
    ]
    
    for cta in cta_text:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(cta)
        set_chinese_font(run, size=11, bold=True, color=RGBColor(26, 86, 219))
        
    # --- Footer Info ---
    doc.add_section()
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(36)
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run("Tech Specs: OpenAI GPT + Vector DB | Security: JWT, Multi-tenant Isolation\nUniAI人資長 © 2026")
    set_chinese_font(run_foot, size=8, color=RGBColor(150, 150, 150))

    output_path = r"c:\Users\User\Desktop\aihr\docs\PRODUCT_INTRODUCTION.docx"
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_document()
